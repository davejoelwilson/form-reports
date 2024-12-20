import os
import sys
from datetime import datetime
import pandas as pd
import pytz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.shared import qn
from connectwise_report.utils.formatting import set_cell_background, format_detail, format_detail_cell

class WordReport:
    def __init__(self):
        self.document = Document()
        self.set_default_styles()

    def set_default_styles(self):
        style = self.document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
    def clean_ticket_summary(self, summary):
        """Clean up ticket summary by removing specific phrases"""
        if not summary:
            return 'N/A'
        
        # List of phrases to remove (including variations with brackets)
        remove_phrases = [
            "(DONT ADD TIME speak with Chris)",
            "[DONT ADD TIME speak with Chris]",
            "(DONT ADD TIME speak with chris)",
            "[DONT ADD TIME speak with chris]",
            "(DONT ADD TIME)",
            "[DONT ADD TIME]",
            "DONT ADD TIME",
            "(speak with Chris)",
            "[speak with Chris]",
            "speak with Chris",
            "(speak with chris)",
            "[speak with chris]",
            "speak with chris"
        ]
        
        cleaned_summary = summary
        for phrase in remove_phrases:
            cleaned_summary = cleaned_summary.replace(phrase, '').strip()
        
        # Clean up any leftover empty brackets
        cleaned_summary = cleaned_summary.replace('()', '').replace('[]', '').strip()
        
        return cleaned_summary or 'N/A'  # Return 'N/A' if empty after cleaning

    def create_entry_table(self, entry):
        # Add some spacing before table
        self.document.add_paragraph()
        
        # Create table with 7 rows (added one for description)
        table = self.document.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        table.allow_autofit = False
        
        # Set table-level properties to prevent auto-fitting with proper namespace
        table._tbl.xpath('./w:tblPr')[0].append(
            parse_xml('''
                <w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                            w:type="fixed"/>
            ''')
        )
        
        # Set absolute widths in twips (1 inch = 1440 twips)
        first_col_width = 1800  # About 1.25 inches (back to original)
        second_col_width = 7920  # About 5.5 inches (adjusted accordingly)
        
        # Apply widths to all cells consistently
        for row in table.rows:
            # First column
            cell = row.cells[0]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(first_col_width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            
            # Second column
            cell = row.cells[1]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(second_col_width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # Define light green color (RGB: 204, 255, 204)
        green_color = "CCFFCC"

        def shorten_url(text):
            """Helper to shorten URLs in text"""
            # Look for SharePoint URLs with specific patterns
            if 'sharepoint.com' in text.lower():
                # Extract the base part of the URL
                if '?sourcedoc=' in text:
                    base_url = text.split('?sourcedoc=')[0]
                    return f"{base_url} [...]"
                elif '&file=' in text:
                    base_url = text.split('&file=')[0]
                    return f"{base_url} [...]"
            return text

        # Convert UTC to NZ time for both start and end times
        start_date = datetime.fromisoformat(entry['timeStart'].replace('Z', '+00:00'))
        end_date = datetime.fromisoformat(entry['timeEnd'].replace('Z', '+00:00'))
        nz_tz = pytz.timezone('Pacific/Auckland')
        nz_start = start_date.astimezone(nz_tz)
        nz_end = end_date.astimezone(nz_tz)

        # Get ticket info
        ticket = entry.get('ticket', {})
        ticket_id = str(ticket.get('id', 'N/A'))
        ticket_summary = self.clean_ticket_summary(ticket.get('summary', 'N/A'))

        # Populate table rows with correct field mappings
        rows = [
            ("Date", nz_start.strftime('%d-%m-%Y')),
            ("Start Time", nz_start.strftime('%I:%M %p')),
            ("End Time", nz_end.strftime('%I:%M %p')),
            ("IT360 Ticket Ref", ticket_id),
            ("Site Name", ticket_summary),  # Added new row
            ("Engineer", entry.get('member', {}).get('name', 'N/A')),
            ("Detail", entry.get('notes', ''))
        ]

        for i, (label, value) in enumerate(rows):
            # Format left column
            left_cell = table.cell(i, 0)
            left_cell.text = label
            left_cell.paragraphs[0].runs[0].font.bold = True
            self.set_cell_background(left_cell, green_color)
            
            # Format right column
            right_cell = table.cell(i, 1)
            if label == "Detail":
                # Handle multiline detail text
                lines = value.split('\n')
                formatted_lines = []
                for line in lines:
                    if line.strip():  # Only process non-empty lines
                        # Shorten URLs in the line
                        line = shorten_url(line.strip())
                        if not line.startswith('-'):
                            line = f"- {line}"
                        formatted_lines.append(line)
                right_cell.text = '\n'.join(formatted_lines)
            else:
                right_cell.text = str(value)

            # Ensure cell widths are fixed
            left_cell._tc.tcPr.tcW.type = 'dxa'
            right_cell._tc.tcPr.tcW.type = 'dxa'

    def set_cell_background(self, cell, color):
        """Helper function to set cell background color"""
        cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>'))

    def generate(self, entries, output_dir, company_id):
        # Add title
        title = self.document.add_paragraph("Weekly Activity Report")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.bold = True
        title.runs[0].font.size = Pt(14)

        # Filter out entries with excluded words in ticket summary
        excluded_words = ["Meetings", "Documentation"]
        filtered_entries = [
            entry for entry in entries 
            if not any(word.lower() in entry.get('ticket', {}).get('summary', '').lower() 
                      for word in excluded_words)
        ]

        # Sort entries by timeStart
        sorted_entries = sorted(filtered_entries, key=lambda x: x['timeStart'])
        
        # Add date range subheading
        if sorted_entries:
            first_entry = sorted_entries[0]
            last_entry = sorted_entries[-1]
            start_date = datetime.fromisoformat(first_entry['timeStart'].replace('Z', '+00:00'))
            end_date = datetime.fromisoformat(last_entry['timeStart'].replace('Z', '+00:00'))
            nz_tz = pytz.timezone('Pacific/Auckland')
            nz_start = start_date.astimezone(nz_tz)
            nz_end = end_date.astimezone(nz_tz)
            
            date_range = self.document.add_paragraph(f"{nz_start.strftime('%d-%m-%Y')} to {nz_end.strftime('%d-%m-%Y')}")
            date_range.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_range.runs[0].font.size = Pt(12)

        # Create a table for each entry
        for entry in sorted_entries:
            self.create_entry_table(entry)

        # Save document
        timestamp = datetime.now().strftime("%Y%m%d")
        filename = f"{output_dir}/activity_report_{company_id}_{timestamp}.docx"
        self.document.save(filename)
        return filename