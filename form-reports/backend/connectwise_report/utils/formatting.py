from docx.shared import Pt
from docx.oxml import parse_xml
import re

def set_cell_background(cell, color):
    """Helper function to set cell background color"""
    shading_elm = parse_xml(
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_detail(notes, project):
    """Format the detail column with project name and bullet points"""
    if isinstance(project, dict):
        project_name = project.get('name', '').split(' - ')[0]
    else:
        project_name = ''
    
    lines = []
    if project_name:
        lines.append(f"Onsite Support - {project_name}")
    
    if notes:
        # Process URLs in notes
        processed_notes = process_urls(notes)
        for line in processed_notes.split('\n'):
            if line.strip():
                if not line.startswith('#'):
                    line = '# ' + line
                lines.append(line)
    
    return '\n'.join(lines)

def process_urls(text):
    """Replace long URLs with more readable text"""
    # SharePoint URL pattern
    sharepoint_pattern = r'https://it360nz\.sharepoint\.com/[^\s\)]+(?=[\s\)])'
    
    def replace_sharepoint_url(match):
        url = match.group(0)
        # Extract filename if present
        filename = re.search(r'file=([^&]+)', url)
        if filename:
            # Clean up filename
            clean_name = filename.group(1).replace('%20', ' ').split('_')[-1].replace('.docx', '')
            return f"[SharePoint Doc: {clean_name}]"
        return "[SharePoint Link]"
    
    # Replace SharePoint URLs
    text = re.sub(sharepoint_pattern, replace_sharepoint_url, text)
    
    # General long URL pattern (for other types of URLs)
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    
    def replace_long_url(match):
        url = match.group(0)
        if len(url) > 50:  # Only replace if URL is longer than 50 chars
            return "[External Link]"
        return url
    
    # Replace other long URLs
    text = re.sub(url_pattern, replace_long_url, text)
    
    return text

def format_detail_cell(cell, text):
    """Format the detail cell with proper styling"""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            paragraph = cell.add_paragraph()
        
        run = paragraph.add_run(line)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0) 