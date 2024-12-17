import os
import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from datetime import datetime
import pandas as pd
import pytz
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import parse_xml
from utils.formatting import set_cell_background, format_detail, format_detail_cell

class WordReport:
    def generate(self, time_entries, output_dir, company_id):
        """Generate the main Word document report"""
        df = pd.DataFrame(time_entries)
        
        # Export raw data for debugging
        debug_csv = os.path.join(output_dir, f"debug_raw_entries_{company_id}_{datetime.now().strftime('%Y%m%d')}.csv")
        df.to_csv(debug_csv, index=False)
        print(f"Debug CSV saved to: {debug_csv}")
        
        # Ensure proper ticket ID extraction and handling
        def get_ticket_id(ticket):
            if isinstance(ticket, dict):
                return str(ticket.get('id', ''))
            return ''
        
        # Add ticket_id column with proper extraction
        df['ticket_id'] = df['ticket'].apply(get_ticket_id)
        
        # Export processed data for debugging
        debug_processed_csv = os.path.join(output_dir, f"debug_processed_entries_{company_id}_{datetime.now().strftime('%Y%m%d')}.csv")
        df.to_csv(debug_processed_csv, index=False)
        print(f"Debug processed CSV saved to: {debug_processed_csv}")
        
        # Convert times to datetime objects for calculation
        df['timeStart'] = pd.to_datetime(df['timeStart'])
        df['timeEnd'] = pd.to_datetime(df['timeEnd'])
        
        # Format data with NZ timezone and calculate hours
        report_data = pd.DataFrame({
            'Ticket': df['ticket_id'],
            'Date': pd.to_datetime(df['timeStart']).dt.tz_convert('Pacific/Auckland').dt.strftime('%A %d/%m/%Y'),
            'Start Time': pd.to_datetime(df['timeStart']).dt.tz_convert('Pacific/Auckland').dt.strftime('%I:%M %p'),
            'End Time': pd.to_datetime(df['timeEnd']).dt.tz_convert('Pacific/Auckland').dt.strftime('%I:%M %p'),
            'Hours': df['actualHours'],
            'Engineer': df['member'].apply(lambda x: x.get('name', '') if isinstance(x, dict) else ''),
            'Detail': df.apply(lambda row: format_detail(row['notes'], row['project']), axis=1)
        }).sort_values(['Ticket', 'Date', 'Start Time'], ascending=[True, False, True])
        
        # Modify ticket references to include summary
        ticket_refs = df['ticket'].apply(lambda x: {
            'id': f"Project Ticket #{x.get('id', '')}" if isinstance(x, dict) else '',
            'summary': x.get('summary', '') if isinstance(x, dict) else ''
        })
        
        return self._create_document(report_data, ticket_refs, output_dir, company_id)
    
    def _create_document(self, report_data, ticket_refs, output_dir, company_id):
        doc = Document()
        
        # Set landscape orientation and margins
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)
        
        # Add title
        title = doc.add_paragraph("iT360 Support Hours for Form Auckland")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.bold = True
        title.runs[0].font.size = Pt(14)
        
        # Get date range from report_data
        dates = pd.to_datetime(report_data['Date'], format='%A %d/%m/%Y')  # Specify the format
        start_date = dates.min()
        end_date = dates.max()
        
        # Add date range
        date_range = doc.add_paragraph()
        date_range.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_range_text = f"Period: {start_date.strftime('%A %d/%m/%Y')} to {end_date.strftime('%A %d/%m/%Y')}"
        date_range_run = date_range.add_run(date_range_text)
        date_range_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Add spacing
        
        # Add "By Ticket" section header with more prominence
        section_header1 = doc.add_paragraph()
        section_header1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = section_header1.add_run("Section 1: Time Entries By Ticket")
        run1.font.bold = True
        run1.font.size = Pt(14)
        doc.add_paragraph()  # Extra spacing
        
        # Create first table (By Ticket)
        self._create_table(doc, report_data, ticket_refs, group_by_ticket=True)
        
        # Force page break
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        
        # Add "By Date" section header with more prominence
        section_header2 = doc.add_paragraph()
        section_header2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = section_header2.add_run("Section 2: Time Entries By Date")
        run2.font.bold = True
        run2.font.size = Pt(14)
        doc.add_paragraph()  # Extra spacing
        
        # Create second table (By Date)
        self._create_table(doc, report_data, ticket_refs, group_by_ticket=False)
        
        # Save document
        output_file = os.path.join(output_dir, f"company_{company_id}_report_{datetime.now().strftime('%Y%m%d')}.docx")
        doc.save(output_file)
        print(f"Report saved to: {output_file}")
        return output_file
    
    def _create_table(self, doc, report_data, ticket_refs, group_by_ticket=True):
        """Create a table grouped either by ticket or by date"""
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Calculate usable width (page width minus margins)
        usable_width = Inches(11.69 - 0.8)  # Page width (11.69) minus margins (0.4 * 2)
        detail_width = Inches((11.69 - 0.8) * 0.5)  # 50% of usable width
        other_width = Inches((11.69 - 0.8) * 0.125)  # Remaining 50% split between 4 columns
        
        # Set column widths
        widths = [
            other_width,     # Date
            other_width,     # Start Time
            other_width,     # End Time
            other_width,     # Engineer
            detail_width     # Detail - 50% of width
        ]
        
        # Apply widths and force them to stick
        for i, width in enumerate(widths):
            table.columns[i].width = width
            # Force width on each cell
            for cell in table.columns[i].cells:
                cell._tc.tcPr.tcW.type = 'dxa'
                cell._tc.tcPr.tcW.w = width.emu
        
        # Add headers with background color
        headers = ['Date', 'Start Time', 'End Time', 'Engineer', 'Detail']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            set_cell_background(cell, "E5F3E2")
        
        if group_by_ticket:
            self._add_data_rows(table, report_data, ticket_refs)
        else:
            # Sort by date and add rows without ticket grouping
            sorted_data = report_data.sort_values(['Date', 'Start Time'])
            current_date = None
            current_date_hours = 0
            
            for _, row in sorted_data.iterrows():
                if current_date != row['Date']:
                    if current_date is not None:
                        self._add_subtotal_row(table, current_date_hours)
                        self._add_separator_row(table)
                    
                    current_date = row['Date']
                    current_date_hours = 0
                    # Add date header
                    date_row = table.add_row().cells
                    date_row[0].merge(date_row[-1])
                    date_row[0].text = current_date
                    set_cell_background(date_row[0], "E5F3E2")
                    
                    for paragraph in date_row[0].paragraphs:
                        paragraph.paragraph_format.space_before = Pt(12)
                        paragraph.paragraph_format.space_after = Pt(12)
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                
                # Add data row
                cells = table.add_row().cells
                cells[0].text = str(row['Date'])
                cells[1].text = str(row['Start Time'])
                cells[2].text = str(row['End Time'])
                cells[3].text = str(row['Engineer'])
                format_detail_cell(cells[4], str(row['Detail']))
                
                # Update hours total
                current_date_hours += float(row['Hours'])
            
            # Add final subtotal
            if current_date is not None:
                self._add_subtotal_row(table, current_date_hours)
        
        # Add grand total
        grand_total = report_data['Hours'].sum()
        doc.add_paragraph()
        total_para = doc.add_paragraph()
        total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_run = total_para.add_run(f"Total Hours: {grand_total:.2f}")
        total_run.font.bold = True
        total_run.font.size = Pt(12)
        total_run.font.name = 'Calibri'
    
    def _add_data_rows(self, table, report_data, ticket_refs):
        current_ticket = None
        current_ticket_hours = 0
        
        for (_, row), ticket_ref in zip(report_data.iterrows(), ticket_refs):
            if current_ticket != ticket_ref['id']:  # Changed to access id from dict
                if current_ticket is not None:
                    self._add_subtotal_row(table, current_ticket_hours)
                    self._add_separator_row(table)
                
                current_ticket = ticket_ref['id']
                current_ticket_hours = 0
                self._add_ticket_header(table, ticket_ref)
            
            # Add data row
            cells = table.add_row().cells
            cells[0].text = str(row['Date'])
            cells[1].text = str(row['Start Time'])
            cells[2].text = str(row['End Time'])
            cells[3].text = str(row['Engineer'])
            format_detail_cell(cells[4], str(row['Detail']))
            
            # Update hours total
            current_ticket_hours += float(row['Hours'])
            
            # Format cells with reduced spacing
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(3)    # Reduced from 6
                    paragraph.paragraph_format.space_before = Pt(3)   # Reduced from 6
                    paragraph.paragraph_format.line_spacing = 1.0     # Reduced from 1.15
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
        
        # Add final subtotal
        if current_ticket is not None:
            self._add_subtotal_row(table, current_ticket_hours)
    
    def _add_subtotal_row(self, table, hours):
        subtotal_row = table.add_row().cells
        subtotal_row[0].merge(subtotal_row[3])
        subtotal_row[0].text = "Subtotal Hours:"
        subtotal_row[4].text = f"{hours:.2f}"
        
        for cell in subtotal_row:
            set_cell_background(cell, "F2F2F2")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    
    def _add_separator_row(self, table):
        separator_cells = table.add_row().cells
        for cell in separator_cells:
            set_cell_background(cell, "F2F2F2")
            cell.text = ""
            # Reduce separator height
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:trHeight w:val="200" w:hRule="atLeast" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
    
    def _add_ticket_header(self, table, ticket_ref):
        ticket_row = table.add_row().cells
        ticket_row[0].merge(ticket_row[-1])
        # Format as "IT360 Project Ticket #XXXXX - Ticket Summary"
        header_text = f"IT360 {ticket_ref['id']}"
        if ticket_ref['summary']:
            header_text += f" - {ticket_ref['summary']}"
        ticket_row[0].text = header_text
        set_cell_background(ticket_row[0], "E5F3E2")
        
        # Reduce padding in ticket headers
        for paragraph in ticket_row[0].paragraphs:
            paragraph.paragraph_format.space_before = Pt(6)   # Reduced from 12
            paragraph.paragraph_format.space_after = Pt(6)    # Reduced from 12
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)